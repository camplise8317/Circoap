import streamlit as st
import google.generativeai as genai
import docx
import io
import openai

# --- Configuración de API Keys (se mantiene igual) ---
st.sidebar.header("Configuración de API Keys")
gemini_api_key = st.sidebar.text_input("API Key de Google Gemini", type="password", 
                                        value=st.session_state.get("gemini_api_key", ""), 
                                        help="Obtén tu clave en https://aistudio.google.com/app/apikey")
openai_api_key = st.sidebar.text_input("API Key de OpenAI (para modelos GPT)", type="password", 
                                       value=st.session_state.get("openai_api_key", ""), 
                                       help="Obtén tu clave en https://platform.openai.com/account/api-keys")

if gemini_api_key:
    st.session_state["gemini_api_key"] = gemini_api_key
if openai_api_key:
    st.session_state["openai_api_key"] = openai_api_key

gemini_config_ok = False
openai_config_ok = False

if gemini_api_key:
    try:
        genai.configure(api_key=gemini_api_key)
        gemini_config_ok = True
        st.sidebar.success("API Key de Gemini configurada.")
    except Exception as e:
        st.sidebar.error(f"Error al configurar la API Key de Gemini: {e}")
else:
    st.sidebar.warning("Por favor, ingresa tu API Key de Gemini.")

if openai_api_key:
    openai.api_key = openai_api_key
    openai_config_ok = True
    st.sidebar.success("API Key de OpenAI configurada.")
else:
    st.sidebar.warning("Por favor, ingresa tu API Key de OpenAI.")

# --- Diccionario Detallado de la Taxonomía de Bloom ---
bloom_taxonomy_detallada = {
    "RECORDAR": { "definicion": "Recuperar conocimiento relevante de la memoria de largo plazo.", "subprocesos": { "Reconocer": {"nombres_alternativos": "Identificar", "definicion_ejemplo": "Localizar conocimiento..."}, "Evocar": {"nombres_alternativos": "Recuperar", "definicion_ejemplo": "Recuperar conocimiento..."} } },
    "COMPRENDER": { "definicion": "Construir significado a partir de contenidos educativos.", "subprocesos": { "Interpretar": {"nombres_alternativos": "Aclarar, parafrasear", "definicion_ejemplo": "Transformar de una forma de representación a otra..."}, "Ejemplificar": {"nombres_alternativos": "Ilustrar, citar casos", "definicion_ejemplo": "Poner un ejemplo específico..."}, "Clasificar": {"nombres_alternativos": "Categorizar", "definicion_ejemplo": "Determinar que algo pertenece a una categoría..."}, "Resumir": {"nombres_alternativos": "Abstraer, generalizar", "definicion_ejemplo": "Extraer el tema general..."}, "Inferir": {"nombres_alternativos": "Concluir, predecir", "definicion_ejemplo": "Sacar una conclusión lógica..."}, "Comparar": {"nombres_alternativos": "Contrastar, esquematizar", "definicion_ejemplo": "Detectar correspondencias..."}, "Explicar": {"nombres_alternativos": "Construir modelos", "definicion_ejemplo": "Construir un modelo de causa-efecto..."} } },
    "APLICAR": { "definicion": "Desarrollar o usar un procedimiento en una situación dada.", "subprocesos": { "Ejecutar": {"nombres_alternativos": "Llevar a cabo", "definicion_ejemplo": "Aplicar un procedimiento a una tarea familiar..."}, "Implementar": {"nombres_alternativos": "Utilizar", "definicion_ejemplo": "Aplicar un procedimiento a una tarea no familiar..."} } },
    "ANALIZAR": { "definicion": "Despiezar el material en sus partes constituyentes y determinar cómo se relacionan.", "subprocesos": { "Diferenciar": {"nombres_alternativos": "Discriminar, seleccionar", "definicion_ejemplo": "Distinguir las partes relevantes..."}, "Organizar": {"nombres_alternativos": "Integrar, estructurar", "definicion_ejemplo": "Determinar cómo encajan los elementos..."}, "Atribuir": {"nombres_alternativos": "Deconstruir", "definicion_ejemplo": "Determinar los puntos de vista, sesgos..."} } },
    "EVALUAR": { "definicion": "Formular juicios con base en criterios o parámetros.", "subprocesos": { "Verificar": {"nombres_alternativos": "Detectar, monitorear", "definicion_ejemplo": "Detectar inconsistencias o falacias..."}, "Criticar": {"nombres_alternativos": "Juzgar, argumentar", "definicion_ejemplo": "Detectar inconsistencias con base en criterios externos..."} } },
    "CREAR": { "definicion": "Agrupar elementos para formar un todo coherente o funcional.", "subprocesos": { "Generar": {"nombres_alternativos": "Formular hipótesis", "definicion_ejemplo": "Formular hipótesis alternativas..."}, "Planear": {"nombres_alternativos": "Diseñar", "definicion_ejemplo": "Idear un procedimiento..."}, "Producir": {"nombres_alternativos": "Construir", "definicion_ejemplo": "Inventar un producto..."} } }
}

# --- Sistema de Prompts Centralizado ---
def get_master_prompt_system():
    bloom_text = ""
    for level, data in bloom_taxonomy_detallada.items():
        bloom_text += f"\n### {level}: {data['definicion']}\n"
        for subprocess, sub_data in data.get('subprocesos', {}).items():
            alt_names = sub_data.get('nombres_alternativos', '')
            bloom_text += f"- **{subprocess} ({alt_names}):** {sub_data.get('definicion_ejemplo', '')}\n"

    MASTER_PROMPT_SYSTEM = f"""
# MODELO PEDAGÓGICO INTEGRAL PARA DISEÑO DE ACTIVIDADES
## CAPA 1: FILOSOFÍA (Círculos de Aprendizaje)
Entorno colaborativo, no competitivo. El facilitador es un guía que usa la mayéutica (preguntas) para fomentar el descubrimiento.
## CAPA 2: ESTRUCTURA (Bruner)
El aprendizaje sigue el viaje: Enactivo (hacer) -> Icónico (representar) -> Simbólico (abstraer).
## CAPA 3: COHESIÓN (Hilo Conductor)
El producto de una fase es el insumo de la siguiente, creando una cadena de evidencia.
## CAPA 4: DIFERENCIACIÓN (Piso Bajo, Techo Alto)
Cada fase debe ser accesible para todos (Piso Bajo) y desafiante para los más avanzados (Techo Alto).
## CAPA 5: INTENCIÓN COGNITIVA (Bloom)
Las tareas deben provocar procesos de pensamiento específicos. El flujo general debe ascender en la taxonomía.
## CAPA 6: DETALLE DE PROCESOS COGNITIVOS (Taxonomía de Bloom Detallada)
Utiliza los siguientes verbos, definiciones y subprocesos para diseñar las preguntas y tareas con la máxima precisión.
{bloom_text}
"""
    return MASTER_PROMPT_SYSTEM

# --- Estructura de categorías (se mantiene igual) ---
CATEGORIAS_ACTIVIDADES = {
    "Círculos de Matemática y Razonamiento": {"Edades": ["5 a 7 años", "8 a 11 años", "12 a 15 años"]},
    "Ciencias": {"Disciplinas": ["Física", "Química", "Biología"]},
    "Tecnología": {"Disciplinas": ["Programación", "Robótica"]}
}

# --- Función para generar texto con LLM (se mantiene igual) ---
def generar_texto_con_llm(model_type, model_name, prompt):
    if model_type == "Gemini":
        if not gemini_config_ok:
            st.error("API Key de Gemini no configurada.")
            return None
        modelo = genai.GenerativeModel(model_name)
        response = modelo.generate_content(prompt)
        return response.text
    elif model_type == "GPT":
        if not openai_config_ok:
            st.error("API Key de OpenAI no configurada.")
            return None
        client = openai.OpenAI(api_key=openai.api_key)
        try:
            response = client.chat.completions.create(
                model=model_name,
                messages=[{"role": "user", "content": prompt}],
                max_tokens=4000
            )
            return response.choices[0].message.content
        except Exception as e:
            st.error(f"Ocurrió un error al llamar a la API de OpenAI: {e}")
            return None
    return None

# --- Función de Auditoría Actualizada ---
def auditar_actividad_circulo_aprendizaje(model_type, model_name, actividad_generada, nivel_salida_esperado):
    MASTER_PROMPT_SYSTEM = get_master_prompt_system()
    
    auditoria_prompt = f"""
    Eres un auditor experto en diseño instruccional. Tu tarea es auditar RIGUROSAMENTE la siguiente actividad.

    --- MODELO PEDAGÓGICO DE REFERENCIA ---
    {MASTER_PROMPT_SYSTEM}
    --------------------------------------
    
    --- OBJETIVO COGNITIVO PARA ESTA ACTIVIDAD ---
    El diseñador solicitó que el nivel de salida esperado de los estudiantes alcance la habilidad de **{nivel_salida_esperado}**.
    --------------------------------------

    --- ACTIVIDAD A AUDITAR ---
    {actividad_generada}
    --------------------

    **VALIDACIÓN DE CRITERIOS (Responde con ✅/❌ y un comentario breve si es ❌):**

    1.  **Hilo Conductor:** ¿El producto de cada fase se usa explícitamente como insumo de la siguiente?
    2.  **Diferenciación:** ¿Cada fase tiene claros puntos de entrada (Piso Bajo) y de profundización (Techo Alto)?
    3.  **Intención Cognitiva (Bloom):** ¿La actividad refleja un ascenso en la taxonomía y culmina exitosamente en el nivel de **{nivel_salida_esperado}** en la fase simbólica?
    4.  **Filosofía del Círculo:** ¿La actividad promueve la colaboración y la discusión?

    **DICTAMEN FINAL:** [✅ CUMPLE / ⚠️ CUMPLE PARCIALMENTE / ❌ RECHAZADO]
    **OBSERVACIONES FINALES:** [Si no es ✅, sé específico en qué capa del modelo falló, especialmente si no se alcanzó el nivel de salida esperado.]
    """
    return generar_texto_con_llm(model_type, model_name, auditoria_prompt)

# --- Función de Generación Actualizada ---
def generar_actividad_circulo_aprendizaje(gen_model_type, gen_model_name, audit_model_type, audit_model_name,
                                          tema_foco_usuario, subcategoria_seleccionada,
                                          nivel_entrada, nivel_salida):
    MASTER_PROMPT_SYSTEM = get_master_prompt_system()
    
    current_activity_text = ""
    auditoria_status = "❌ RECHAZADO"
    audit_observations = ""
    max_refinement_attempts = 3
    attempt = 0
    activity_final_data = None
    
    while auditoria_status != "✅ CUMPLE" and attempt < max_refinement_attempts:
        attempt += 1
        st.info(f"--- Generando/Refinando Actividad (Intento {attempt}/{max_refinement_attempts}) ---")

        prompt_content_for_llm = f"""
        Eres un diseñador instruccional de élite. Tu tarea es crear una actividad de 1 hora.
        Debes seguir de forma ESTRICTA Y DETALLADA el modelo pedagógico integral proporcionado.

        --- 1. ENTRADA ESTRATÉGICA (DEFINICIÓN DEL RETO) ---
        - **Tema:** {tema_foco_usuario}
        - **Grupo:** {subcategoria_seleccionada}
        - **Nivel de Entrada Esperado:** {nivel_entrada} (Diseña el 'Piso Bajo' de la fase enactiva para este nivel).
        - **Nivel de Salida Esperado:** {nivel_salida} (El 'Techo Alto' de la fase simbólica debe culminar en este nivel de Bloom).

        --- 2. MODELO PEDAGÓGICO A IMPLEMENTAR ---
        {MASTER_PROMPT_SYSTEM}
        -------------------------------------------

        --- 3. FORMATO ESTRICTO DE SALIDA (Aplica el Modelo) ---
        **TÍTULO DE LA ACTIVIDAD:** [Título creativo y atractivo]

        **OBJETIVOS DE APRENDIZAJE (CON VERBOS DE BLOOM):**
        - [Define 2-3 objetivos que culminen en el nivel de salida esperado ({nivel_salida}). Usa verbos específicos de la Capa 6 del modelo.]

        **EL HILO CONDUCTOR (LA CADENA DE EVIDENCIA):**
        - **Artefacto Enactivo:** [Define el producto físico que se creará.]
        - **Representación Icónica:** [Define el producto visual que lo analizará.]
        - **Conclusión Simbólica:** [Define el producto abstracto que lo generalizará.]

        **DESARROLLO DE LA ACTIVIDAD (60 MINUTOS)**
        ---
        **FASE 1: ENACTIVA (20 min) | Foco Cognitivo: APLICAR**
        - **Facilitador (Piso Bajo):** [Describe la invitación abierta a la exploración, alineada con el Nivel de Entrada.]
        - **Facilitador (Techo Alto):** [Describe 1-2 desafíos de profundización que eleven la cognición hacia ANALIZAR.]
        - **➡️ Producto Clave (Insumo para Fase 2):** [Confirma el artefacto enactivo.]
        ---
        **FASE 2: ICÓNICA (20 min) | Foco Cognitivo: ANALIZAR**
        - **Punto de Partida:** El Artefacto Enactivo.
        - **Facilitador (Piso Bajo):** [Pregunta para representar. Usa verbos de COMPRENDER-Interpretar/Explicar.]
        - **Facilitador (Techo Alto):** [Reto de sistematización. Usa verbos de ANALIZAR-Organizar o CREAR-Planear.]
        - **➡️ Producto Clave (Insumo para Fase 3):** [Confirma la representación icónica.]
        ---
        **FASE 3: SIMBÓLICA (15 min) | Foco Cognitivo: {nivel_salida}**
        - **Punto de Partida:** La Representación Icónica.
        - **Facilitador (Piso Bajo):** [Pregunta para explicar y comparar. Usa verbos de ANALIZAR-Comparar.]
        - **Facilitador (Techo Alto):** [Pregunta para juzgar y generalizar, usando explícitamente verbos del nivel de salida ({nivel_salida}) de la Capa 6.]
        - **➡️ Producto Clave (Resultado Final):** [Confirma la conclusión simbólica, que debe reflejar el nivel de salida.]
        ---
        **CIERRE Y REFLEXIÓN (5 min):**
        - [Resume el viaje del Hilo Conductor, desde el objeto hasta la idea.]
        """
        
        if attempt > 1:
            prompt_content_for_llm += f"\n--- RETROALIMENTACIÓN PARA REFINAMIENTO ---\nLa versión anterior fue rechazada. Observaciones del auditor: {audit_observations}\nPor favor, genera una nueva versión que corrija estos puntos.\n"

        full_llm_response = generar_texto_con_llm(gen_model_type, gen_model_name, prompt_content_for_llm)

        if not full_llm_response:
            st.error("Fallo en la generación de texto. El intento se detendrá.")
            break
        
        current_activity_text = full_llm_response
        st.subheader(f"Actividad Generada (Intento {attempt}):")
        st.markdown(current_activity_text)
        
        auditoria_resultado = auditar_actividad_circulo_aprendizaje(audit_model_type, audit_model_name, current_activity_text, nivel_salida)

        if not auditoria_resultado:
            st.error("Fallo en la auditoría. El intento se detendrá.")
            break
            
        st.subheader(f"Resultado de Auditoría (Intento {attempt}):")
        st.markdown(auditoria_resultado)
        
        if "✅ CUMPLE" in auditoria_resultado:
            auditoria_status = "✅ CUMPLE"
            st.success(f"¡Actividad generada y aprobada en el intento {attempt}!")
            break
        else:
            observaciones_start = auditoria_resultado.find("OBSERVACIONES FINALES:")
            if observaciones_start != -1:
                audit_observations = auditoria_resultado[observaciones_start:].strip()
            else:
                audit_observations = "No se pudieron extraer observaciones."
            st.warning("La actividad necesita refinamiento. Intentando de nuevo...")

    activity_final_data = {
        "activity_text": current_activity_text,
        "classification": {
            "Tema de Foco": tema_foco_usuario,
            "Grupo": subcategoria_seleccionada,
            "Nivel de Entrada": nivel_entrada,
            "Nivel de Salida": nivel_salida
        },
        "final_audit_status": auditoria_status,
        "final_audit_observations": audit_observations
    }
    
    if auditoria_status != "✅ CUMPLE":
        st.error(f"No se pudo generar una actividad aprobada después de {max_refinement_attempts} intentos.")

    return [activity_final_data]

# --- Función para exportar a Word (revisa que esté completa) ---
def exportar_actividad_a_word(actividades_procesadas_list):
    doc = docx.Document()
    doc.add_heading('Actividad de Círculo de Aprendizaje Generada', level=1)
    
    if not actividades_procesadas_list:
        doc.add_paragraph('No se procesó ninguna actividad.')
    else:
        activity_data = actividades_procesadas_list[0]
        classification = activity_data.get("classification", {})
        
        doc.add_heading('Definición Estratégica', level=2)
        doc.add_paragraph(f"**Tema:** {classification.get('Tema de Foco', 'N/A')}")
        doc.add_paragraph(f"**Grupo:** {classification.get('Grupo', 'N/A')}")
        doc.add_paragraph(f"**Nivel de Entrada:** {classification.get('Nivel de Entrada', 'N/A')}")
        doc.add_paragraph(f"**Nivel de Salida:** {classification.get('Nivel de Salida', 'N/A')}")
        doc.add_paragraph('')

        doc.add_heading('Actividad Generada', level=2)
        # Formateo mejorado para el texto de la actividad
        lines = activity_data.get("activity_text", "").split('\n')
        for line in lines:
            stripped_line = line.strip()
            if not stripped_line:
                continue
            
            # Encabezados principales
            if stripped_line.startswith("**TÍTULO") or stripped_line.startswith("**OBJETIVOS") or \
               stripped_line.startswith("**EL HILO") or stripped_line.startswith("**DESARROLLO") or \
               stripped_line.startswith("**CIERRE"):
                p = doc.add_paragraph()
                run = p.add_run(stripped_line.replace("**", ""))
                run.bold = True
                run.font.size = docx.shared.Pt(14)
            # Encabezados de Fases
            elif stripped_line.startswith("---"):
                doc.add_paragraph('---')
            elif stripped_line.startswith("**FASE"):
                p = doc.add_paragraph()
                run = p.add_run(stripped_line.replace("**", ""))
                run.bold = True
                run.font.size = docx.shared.Pt(12)
            # Sub-encabezados de las fases
            elif stripped_line.startswith("- **Facilitador") or stripped_line.startswith("- **➡️ Producto Clave") or \
                 stripped_line.startswith("- **Punto de Partida"):
                p = doc.add_paragraph()
                run = p.add_run(stripped_line.replace("- **", "").replace("**", ""))
                run.bold = True
            else:
                doc.add_paragraph(stripped_line.replace("*", "").strip())

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# --- Interfaz de Usuario de Streamlit ---
st.title("📚 Generador de Actividades Pedagógicas con IA 🧠")
st.markdown("Esta herramienta diseña actividades siguiendo un modelo pedagógico integral (Círculos, Bruner, Bloom).")

# --- Selección de Modelos ---
st.sidebar.header("Configuración de Modelos de IA")
st.sidebar.subheader("Modelo para Generación")
gen_model_type = st.sidebar.radio("Tipo", ["GPT", "Gemini"], key="gen_type")
if gen_model_type == "GPT":
    gen_model_name = st.sidebar.selectbox("Modelo GPT", ["gpt-4o", "gpt-4-turbo", "gpt-3.5-turbo"], key="gen_gpt_name")
else:
    gen_model_name = st.sidebar.selectbox("Modelo Gemini", ["gemini-1.5-pro", "gemini-1.5-flash"], key="gen_gemini_name")

st.sidebar.subheader("Modelo para Auditoría")
audit_model_type = st.sidebar.radio("Tipo", ["Gemini", "GPT"], key="audit_type", index=0) # Gemini por defecto
if audit_model_type == "Gemini":
    audit_model_name = st.sidebar.selectbox("Modelo Gemini", ["gemini-1.5-pro", "gemini-1.5-flash"], key="audit_gemini_name")
else:
    audit_model_name = st.sidebar.selectbox("Modelo GPT", ["gpt-4o", "gpt-4-turbo", "gpt-3.5-turbo"], key="audit_gpt_name")


if gemini_config_ok or openai_config_ok:
    st.header("1. Defina la Entrada Estratégica")

    tema_foco_usuario = st.text_input(
        "**¿QUÉ VAMOS A APRENDER?** (Tema central de la actividad)",
        placeholder="Ej: Principios de estabilidad y centro de gravedad"
    )

    categoria_seleccionada = st.selectbox(
        "Categoría de la actividad",
        list(CATEGORIAS_ACTIVIDADES.keys())
    )
    
    if categoria_seleccionada == "Círculos de Matemática y Razonamiento":
        sub_options = CATEGORIAS_ACTIVIDADES[categoria_seleccionada]["Edades"]
        subcategoria_label = "Grupo (Rango de Edad)"
    else:
        sub_options = CATEGORIAS_ACTIVIDADES[categoria_seleccionada]["Disciplinas"]
        subcategoria_label = "Grupo (Disciplina)"
    
    subcategoria_seleccionada = st.selectbox(subcategoria_label, sub_options)

    st.header("2. Defina el Alcance del Aprendizaje")
    
    nivel_entrada_usuario = st.text_input(
        "**¿DESDE DÓNDE PARTIMOS?** (Nivel de entrada de los estudiantes)",
        placeholder="Ej: Saben construir con bloques, pero sin criterio claro."
    )
    
    nivel_salida_usuario = st.selectbox(
        "**¿A DÓNDE QUEREMOS LLEGAR?** (Máxima habilidad de Bloom a alcanzar)",
        options=list(bloom_taxonomy_detallada.keys()),
        index=len(bloom_taxonomy_detallada) - 1, # Por defecto, 'CREAR'
        help="Esta será la habilidad cognitiva principal de la fase simbólica."
    )

    if st.button("Generar y Auditar Actividad"):
        if not tema_foco_usuario or not nivel_entrada_usuario:
            st.error("Por favor, complete todos los campos de la entrada estratégica.")
        else:
            st.markdown("---")
            st.info("Iniciando ciclo de generación y auditoría...")
            
            # --- CORRECCIÓN CLAVE AQUÍ ---
            # Usamos las variables gen_model_name y audit_model_name de la UI
            activity_processed_list = generar_actividad_circulo_aprendizaje(
                gen_model_type, gen_model_name,
                audit_model_type, audit_model_name,
                tema_foco_usuario,
                subcategoria_seleccionada,
                nivel_entrada_usuario,
                nivel_salida_usuario
            )

            if activity_processed_list:
                st.session_state['last_processed_activity_data'] = activity_processed_list[0]
                final_data = st.session_state['last_processed_activity_data']
                st.subheader("Resultado Final del Proceso")
                
                if final_data['activity_text']:
                    st.markdown(final_data['activity_text'])
                    st.info(f"**Dictamen Final del Auditor:** {final_data['final_audit_status']}")
                    if final_data['final_audit_status'] != "✅ CUMPLE":
                        st.warning(f"**Observaciones del Auditor:** {final_data['final_audit_observations']}")
                else:
                    st.error("No se pudo generar el texto de la actividad. Revisa los logs o la configuración de la API.")

    st.header("3. Exportar Actividad")
    if 'last_processed_activity_data' in st.session_state and st.session_state['last_processed_activity_data']:
        st.success("Hay una actividad lista para ser exportada.")
        nombre_base = tema_foco_usuario.strip().replace(' ', '_')[:30]
        
        word_buffer = exportar_actividad_a_word([st.session_state['last_processed_activity_data']])
        st.download_button(
            label="Descargar Actividad en Word",
            data=word_buffer,
            file_name=f"actividad_{nombre_base}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    else:
        st.info("Genere una actividad para poder exportarla.")

else:
    st.info("Por favor, ingresa al menos una API Key de Gemini o OpenAI en la barra lateral para comenzar.")
