import streamlit as st
import vertexai
from vertexai.generative_models import GenerativeModel
import docx
import io
import os
import re # Importado para ayudar a separar las guías

# --- INICIALIZACIÓN DEL ESTADO DE LA SESIÓN ---
if 'stage' not in st.session_state:
    st.session_state.stage = "inspiration"
if 'inspiration_text' not in st.session_state:
    st.session_state.inspiration_text = ""
if 'num_actividades' not in st.session_state:
    st.session_state.num_actividades = 1
if 'sequence_plan' not in st.session_state:
    st.session_state.sequence_plan = None
if 'final_context' not in st.session_state:
    st.session_state.final_context = ""
if 'processed_sequence' not in st.session_state:
    st.session_state.processed_sequence = []


# --- FUNCIÓN PRINCIPAL QUE ENVUELVE LA APP ---
def main():
    # --- CONFIGURACIÓN DE LA PÁGINA DE STREAMLIT ---
    st.set_page_config(
        page_title="Orquestador de Secuencias Pedagógicas con Vertex AI",
        page_icon="🎼",
        layout="wide"
    )
    st.title("🎼 Orquestador de Secuencias Pedagógicas con IA 🧠")
    st.markdown("Un co-piloto para diseñar unidades didácticas completas, coherentes e inmersivas.")

    # --- INICIALIZACIÓN Y CONFIGURACIÓN DE VERTEX AI ---
    st.sidebar.header("Configuración de Vertex AI")
    try:
        GCP_PROJECT_ID = os.environ.get("GCP_PROJECT")
        GCP_LOCATION = os.environ.get("GCP_LOCATION")

        if not GCP_PROJECT_ID or not GCP_LOCATION:
            st.sidebar.error("Variables de entorno GCP_PROJECT y GCP_LOCATION no encontradas.")
            st.error("Configura tus variables de entorno de Google Cloud para continuar.")
            st.stop()

        vertexai.init(project=GCP_PROJECT_ID, location=GCP_LOCATION)
        st.sidebar.success(f"✅ Conectado a Vertex AI\nProyecto: {GCP_PROJECT_ID}")
    except Exception as e:
        st.sidebar.error(f"Error al inicializar Vertex AI: {e}")
        st.error("No se pudo conectar con Vertex AI. Verifica la configuración del proyecto y tu autenticación.")
        st.stop()
    
    # --- BLOQUE DE CONFIGURACIÓN DE MODELOS EN LA BARRA LATERAL ---
    st.sidebar.subheader("Selección de Modelos")
    vertex_ai_models = [
        "gemini-2.5-pro",
        "gemini-2.5-flash",
        "gemini-2.5-flash-lite"
    ]
    st.session_state.gen_model_name = st.sidebar.selectbox(
        "**Modelo para Generación y Planificación**",
        vertex_ai_models, index=1, key="gen_vertex_name_sidebar"
    )
    st.session_state.audit_model_name = st.sidebar.selectbox(
        "**Modelo para Auditoría**",
        vertex_ai_models, index=0, key="audit_vertex_name_sidebar",
        help="Se recomienda un modelo potente (ej. Pro) para la auditoría."
    )
    
    # --- DICCIONARIO DE LA TAXONOMÍA DE BLOOM (sin cambios) ---
    bloom_taxonomy_detallada = {
        "RECORDAR": { "definicion": "Recuperar conocimiento relevante de la memoria de largo plazo.", "subprocesos": { "Reconocer": {"nombres_alternativos": "Identificar", "definicion_ejemplo": "Localizar conocimiento..."}, "Evocar": {"nombres_alternativos": "Recuperar", "definicion_ejemplo": "Recuperar conocimiento..."} } },
        "COMPRENDER": { "definicion": "Construir significado a partir de contenidos educativos.", "subprocesos": { "Interpretar": {"nombres_alternativos": "Aclarar, parafrasear", "definicion_ejemplo": "Transformar de una forma de representación a otra..."}, "Ejemplificar": {"nombres_alternativos": "Ilustrar, citar casos", "definicion_ejemplo": "Poner un ejemplo específico..."}, "Clasificar": {"nombres_alternativos": "Categorizar", "definicion_ejemplo": "Determinar que algo pertenece a una categoría..."}, "Resumir": {"nombres_alternativos": "Abstraer, generalizar", "definicion_ejemplo": "Extraer el tema general..."}, "Inferir": {"nombres_alternativos": "Concluir, predecir", "definicion_ejemplo": "Sacar una conclusión lógica..."}, "Comparar": {"nombres_alternativos": "Contrastar, esquematizar", "definicion_ejemplo": "Detectar correspondencias..."}, "Explicar": {"nombres_alternativos": "Construir modelos", "definicion_ejemplo": "Construir un modelo de causa-efecto..."} } },
        "APLICAR": { "definicion": "Desarrolar o usar un procedimiento en una situación dada.", "subprocesos": { "Ejecutar": {"nombres_alternativos": "Llevar a cabo", "definicion_ejemplo": "Aplicar un procedimiento a una tarea familiar..."}, "Implementar": {"nombres_alternativos": "Utilizar", "definicion_ejemplo": "Aplicar un procedimiento a una tarea no familiar..."} } },
        "ANALIZAR": { "definicion": "Despiezar el material en sus partes constituyentes y determinar cómo se relacionan.", "subprocesos": { "Diferenciar": {"nombres_alternativos": "Discriminar, seleccionar", "definicion_ejemplo": "Distinguir las partes relevantes..."}, "Organizar": {"nombres_alternativos": "Integrar, estructurar", "definicion_ejemplo": "Determinar cómo encajan los elementos..."}, "Atribuir": {"nombres_alternativos": "Deconstruir", "definicion_ejemplo": "Determinar los puntos de vista, sesgos..."} } },
        "EVALUAR": { "definicion": "Formular juicios con base en criterios o parámetros.", "subprocesos": { "Verificar": {"nombres_alternativos": "Detectar, monitorear", "definicion_ejemplo": "Detectar inconsistencias o falacias..."}, "Criticar": {"nombres_alternativos": "Juzgar, argumentar", "definicion_ejemplo": "Detectar inconsistencias con base en criterios externos..."} } },
        "CREAR": { "definicion": "Agrupar elementos para formar un todo coherente o funcional.", "subprocesos": { "Generar": {"nombres_alternativos": "Formular hipótesis", "definicion_ejemplo": "Formular hipótesis alternativas..."}, "Planear": {"nombres_alternativos": "Diseñar", "definicion_ejemplo": "Idear un procedimiento..."}, "Producir": {"nombres_alternativos": "Construir", "definicion_ejemplo": "Inventar un producto..."} } }
    }


    # --- SISTEMA DE PROMPTS CENTRALIZADO ---
    def get_master_prompt_system(contexto_narrativo):
        bloom_text = ""
        for level, data in bloom_taxonomy_detallada.items():
            bloom_text += f"\n### {level}: {data['definicion']}\n"
            for subprocess, sub_data in data.get('subprocesos', {}).items():
                alt_names = sub_data.get('nombres_alternativos', '')
                bloom_text += f"- **{subprocess} ({alt_names}):** {sub_data.get('definicion_ejemplo', '')}\n"

        return f"""
# MODELO PEDAGÓGICO INTEGRAL
## CAPA 0: CONTEXTO NARRATIVO
Toda la actividad debe estar inmersa en esta historia:
---
{contexto_narrativo}
---
## CAPA 1: FILOSOFÍA (Círculos de Aprendizaje)
Entorno colaborativo. El facilitador guía con preguntas.
## CAPA 2: ESTRUCTURA (Bruner)
Viaje: Enactivo (hacer) -> Icónico (representar) -> Simbólico (abstraer).
## CAPA 3: COHESIÓN (Hilo Conductor)
El producto de una fase es el insumo de la siguiente.
## CAPA 4: DIFERENCIACIÓN (Piso Bajo, Techo Alto)
Accesible para todos, desafiante para los más avanzados.
## CAPA 5: INTENCIÓN COGNITIVA (Bloom)
Las tareas deben provocar procesos de pensamiento específicos y ascender en la taxonomía.
## CAPA 6: DETALLE DE PROCESOS COGNITIVOS
Usa los siguientes verbos y definiciones con precisión.
{bloom_text}
"""

    # --- Estructura de categorías (sin cambios) ---
    CATEGORIAS_ACTIVIDADES = {
        "Círculos de Matemática y Razonamiento": {"Edades": ["5 a 7 años", "8 a 11 años", "12 a 15 años"]},
        "Ciencias": {"Disciplinas": ["Física", "Química", "Biología"]},
        "Tecnología": {"Disciplinas": ["Programación", "Robótica"]}
    }

    # --- FUNCIONES DE UTILIDAD Y LÓGICA DE LA APP ---

    def leer_docx(file):
        try:
            doc = docx.Document(io.BytesIO(file.read()))
            return '\n'.join([para.text for para in doc.paragraphs])
        except Exception as e:
            st.error(f"Error al leer el archivo Word: {e}")
            return ""

    def generar_texto_con_vertex(model_name, prompt):
        try:
            modelo = GenerativeModel(model_name)
            response = modelo.generate_content(prompt)
            return response.text
        except Exception as e:
            st.error(f"Ocurrió un error al llamar al modelo {model_name} en Vertex AI: {e}")
            return None

    def set_stage(stage_name):
        st.session_state.stage = stage_name

    # --- NUEVA FUNCIÓN "CEREBRO" PARA PLANIFICAR LA SECUENCIA ---
    def planificar_secuencia(inspiration_text, num_actividades, nivel_salida_final, model_name):
        st.info(f"Diseñando un plan de vuelo para {num_actividades} sesiones...")
        prompt_planificacion = f"""
        Eres un experto en diseño curricular. Basado en la siguiente inspiración y requisitos, crea un plan de secuencia de aprendizaje.

        **Inspiración Inicial:** {inspiration_text}
        **Número de Sesiones:** {num_actividades}
        **Nivel Cognitivo Final Deseado (Bloom):** {nivel_salida_final}

        **Tu Tarea:**
        1.  **Define un Hilo Conductor Narrativo:** Crea una historia o misión global que conecte todas las sesiones.
        2.  **Secuencia los Objetivos Cognitivos:** Distribuye los niveles de la Taxonomía de Bloom a lo largo de las {num_actividades} sesiones. Empieza con niveles bajos (Recordar, Comprender) y progresa hacia el nivel final ({nivel_salida_final}). Sé explícito sobre qué nivel de Bloom es el foco principal de cada sesión.
        3.  **Desglosa los Contenidos:** Para cada sesión, define brevemente el sub-tema o concepto específico que se abordará.

        **Formato de Salida (Usa Markdown):**
        
        ### Plan de Secuencia de Aprendizaje
        
        **Hilo Conductor Narrativo:** [Describe la historia o misión global aquí.]
        
        ---
        
        **Sesión 1: [Título de la Sesión 1]**
        - **Concepto Clave:** [Describe el concepto de esta sesión.]
        - **Objetivo Cognitivo (Bloom):** COMPRENDER
        
        **Sesión 2: [Título de la Sesión 2]**
        - **Concepto Clave:** [Describe el concepto de esta sesión.]
        - **Objetivo Cognitivo (Bloom):** APLICAR
        
        ... (continúa para todas las sesiones hasta la {num_actividades})
        """
        plan = generar_texto_con_vertex(model_name, prompt_planificacion)
        return plan

    # --- FUNCIÓN DE AUDITORÍA (sin cambios en su lógica interna) ---
    def auditar_actividad(actividad_generada, nivel_salida_esperado, contexto_narrativo, audit_model_name):
        master_prompt_ref = get_master_prompt_system(contexto_narrativo)
        auditoria_prompt = f"""
        Eres un auditor experto en diseño instruccional. Audita RIGUROSAMENTE la siguiente actividad individual.
        --- MODELO PEDAGÓGICO DE REFERENCIA ---
        {master_prompt_ref}
        --- OBJETIVO COGNITIVO PARA ESTA SESIÓN ---
        El nivel de salida esperado es **{nivel_salida_esperado}**.
        --- ACTIVIDAD A AUDITAR ---
        {actividad_generada}
        ---
        **VALIDACIÓN DE CRITERIOS (Responde con ✅/❌ y un comentario breve si es ❌):**
        1.  **Contexto Narrativo (Capa 0):** ¿La actividad está completamente inmersa en la historia y usa su lenguaje?
        2.  **Hilo Conductor (Capa 3):** ¿El producto de cada fase se usa explícitamente como insumo de la siguiente?
        3.  **Intención Cognitiva (Capa 5):** ¿La actividad culmina exitosamente en el nivel de **{nivel_salida_esperado}** en la fase simbólica?
        **DICTAMEN FINAL:** [✅ CUMPLE / ❌ RECHAZADO]
        **OBSERVACIONES FINALES:** [Si es ❌, sé específico en qué capa del modelo falló.]
        """
        return generar_texto_con_vertex(audit_model_name, auditoria_prompt)

    # --- FUNCIÓN DE GENERACIÓN (PROMPT ACTUALIZADO) ---
    def generar_actividad_con_auditoria(params):
        # El contexto narrativo ahora es el plan de secuencia completo
        master_prompt = get_master_prompt_system(params["plan_secuencia"]) 
        current_activity_text = ""
        audit_observations = ""
        max_attempts = 3
        attempt = 0

        gen_model = params["gen_model"]
        audit_model = params["audit_model"]

        while attempt < max_attempts:
            attempt += 1
            st.info(f"--- Generando/Refinando Sesión {params['session_num']} (Intento {attempt}/{max_attempts}) ---")
            
            # PROMPT ACTUALIZADO PARA GENERAR DOBLE GUÍA
            prompt_generacion = f"""
            Eres un diseñador instruccional de élite. Tu tarea es generar UNA ÚNICA actividad detallada que forma parte de una secuencia mayor.

            ---
            **PLAN DE SECUENCIA GLOBAL (CONTEXTO GENERAL):**
            {params["plan_secuencia"]}
            ---
            **TAREA ESPECÍFICA: Generar la Sesión número {params["session_num"]}**
            ---
            - **Grupo:** {params["grupo"]}
            - **Nivel de Entrada para esta sesión:** {params["nivel_entrada"]}
            - **Modelo Pedagógico Base:** {master_prompt}
            ---

            **FORMATO ESTRICTO DE SALIDA (Genera AMBOS documentos usando Markdown):**

            ---
            ### GUÍA RÁPIDA (PARA EL AULA / FICHA)
            - **Sesión:** {params["session_num"]}
            - **Título:** [Título Atractivo]
            - **Propósito (1-2 líneas):** [Resumen muy breve del objetivo de la sesión.]
            - **Materiales Esenciales:** [Lista simple.]
            - **Momentos Clave (Tiempos Aprox.):**
                - **Momento Enactivo (Hacer):** [Descripción breve de la actividad principal.] (20 min)
                - **Momento Icónico (Representar):** [Descripción breve de cómo se visualizará.] (20 min)
                - **Momento Simbólico (Abstraer):** [Descripción breve de la formalización.] (15 min)
                - **Cierre (Reflexionar):** [Descripción breve.] (5 min)
            ---
            ### GUÍA PARA EL DOCENTE (ACOMPAÑAMIENTO)

            **1. Descripción Detallada:**
               - **Propósito Pedagógico:** [Explicación detallada del porqué de esta actividad.]
               - **Pasos por Fase:**
                 - **Enactiva:** [Instrucciones detalladas, preguntas del facilitador, variantes piso-medio-techo.]
                 - **Icónica:** [Instrucciones detalladas, preguntas, variantes.]
                 - **Simbólica:** [Instrucciones detalladas, preguntas, variantes.]
               - **Cierre:** [Instrucciones para guiar la síntesis, metacognición y próximos pasos.]

            **2. Evaluación Formativa:**
               - **Evidencias a Observar:** [Qué deben producir o decir los estudiantes en cada fase como prueba de comprensión.]
               - **Logros (Criterios de Desempeño):** [Cómo saber si el grupo alcanzó el objetivo cognitivo de la sesión.]
               - **Errores Típicos y Microintervenciones:** [Lista de 2-3 errores comunes y cómo el docente puede intervenir sutilmente.]

            **3. Cohesión y Metacognición:**
               - **Bitácora de Secuencia:** [Cómo esta actividad conecta con la sesión ANTERIOR y prepara la SIGUIENTE.]
               - **Prompts de Metacognición:** [2-3 preguntas específicas para que los estudiantes reflexionen sobre su proceso de aprendizaje al final.]

            **4. Herramientas de Evaluación:**
               - **Rúbrica Analítica Simple:** [Tabla con 2-3 criterios y descriptores observables para las fases clave (ej. Enactiva y Simbólica).]
            """

            if attempt > 1:
                prompt_generacion += f"\n--- RETROALIMENTACIÓN PARA REFINAMIENTO ---\nLa versión anterior fue rechazada. Observaciones del auditor: {audit_observations}\nPor favor, genera una nueva versión que corrija estos puntos.\n"

            current_activity_text = generar_texto_con_vertex(gen_model, prompt_generacion)
            if not current_activity_text:
                st.error("Fallo en la generación de texto.")
                break

            with st.expander(f"Ver Actividad Generada - Sesión {params['session_num']} (Intento {attempt})", expanded=False):
                st.markdown(current_activity_text)
            
            # La auditoría ahora usa el plan como contexto narrativo
            auditoria_resultado = auditar_actividad(current_activity_text, params["nivel_salida"], params["plan_secuencia"], audit_model)
            if not auditoria_resultado:
                st.error("Fallo en la auditoría.")
                break

            with st.expander(f"Ver Resultado de Auditoría - Sesión {params['session_num']} (Intento {attempt})", expanded=True):
                st.markdown(auditoria_resultado)

            if "✅ CUMPLE" in auditoria_resultado:
                st.success(f"¡Sesión {params['session_num']} generada y aprobada en el intento {attempt}!")
                # Extraemos el título para mostrarlo en el expander final
                title_search = re.search(r"\*\*Título:\s*(.*)", current_activity_text)
                title = title_search.group(1) if title_search else f"Sesión {params['session_num']}"
                return {"activity_text": current_activity_text, "status": "✅ CUMPLE", "title": title}
            else:
                observaciones_start = auditoria_resultado.find("OBSERVACIONES FINALES:")
                audit_observations = auditoria_resultado[observaciones_start:] if observaciones_start != -1 else "No se pudo extraer observaciones."
                st.warning(f"La Sesión {params['session_num']} necesita refinamiento...")
        
        st.error(f"No se pudo generar una actividad aprobada para la Sesión {params['session_num']} después de {max_attempts} intentos.")
        return {"activity_text": current_activity_text, "status": "❌ RECHAZADO", "title": f"Sesión {params['session_num']} (Fallida)"}

    # --- FUNCIÓN DE EXPORTACIÓN A WORD (ACTUALIZADA PARA SECUENCIAS) ---
    def exportar_secuencia_a_word(sequence_data):
        doc = docx.Document()
        doc.add_heading('Secuencia de Aprendizaje Generada con IA', level=1)
        
        # Opcional: Añadir el plan de secuencia al inicio del documento
        if st.session_state.sequence_plan:
            doc.add_heading('Plan de Vuelo de la Secuencia', level=2)
            doc.add_paragraph(st.session_state.sequence_plan)
            doc.add_page_break()

        for i, activity_data in enumerate(sequence_data):
            doc.add_heading(f"Actividad de la Sesión {i+1}", level=2)
            
            activity_text = activity_data.get("activity_text", "")
            
            # Lógica para separar las guías (simple split)
            parts = re.split(r'---+\s*### GUÍA PARA EL DOCENTE', activity_text, flags=re.IGNORECASE)
            guia_rapida_text = parts[0].replace("### GUÍA RÁPIDA (PARA EL AULA / FICHA)", "").strip()
            guia_docente_text = ""
            if len(parts) > 1:
                guia_docente_text = parts[1].strip()

            doc.add_heading('Guía Rápida (Ficha de Aula)', level=3)
            doc.add_paragraph(guia_rapida_text)
            
            doc.add_heading('Guía para el Docente (Acompañamiento)', level=3)
            doc.add_paragraph(guia_docente_text)
            
            if i < len(sequence_data) - 1:
                doc.add_page_break()
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    # --- INTERFAZ DE USUARIO POR ETAPAS ---

    if st.session_state.stage == "inspiration":
        st.header("ETAPA 1: El Punto de Partida 💡")
        st.markdown("Define el tema central y la longitud de tu secuencia de aprendizaje.")
        
        tab1, tab2, tab3 = st.tabs(["🎯 Empezar con un Tema", "📝 Lluvia de Ideas", "📄 Subir un Archivo (.docx)"])
        
        def handle_inspiration_submit(inspiration_text):
            if inspiration_text:
                st.session_state.inspiration_text = inspiration_text
                st.session_state.num_actividades = st.session_state.get('num_act_input', 1)
                set_stage("planning")
                st.rerun()
            else:
                st.warning("La fuente de inspiración no puede estar vacía.")

        with tab1:
            tema_foco_usuario = st.text_input("Tema central de la secuencia", placeholder="Ej: El ciclo del agua")
            num_act_input_1 = st.number_input("Número de actividades en la secuencia", min_value=1, max_value=10, value=3, step=1, key="num_act_input_1")
            if st.button("Definir Tema y Continuar", key="tema_btn"):
                st.session_state.num_actividades = num_act_input_1
                handle_inspiration_submit(f"El tema central es: {tema_foco_usuario}.")
        with tab2:
            idea_box = st.text_area("Lluvia de ideas para la secuencia...", height=200, placeholder="Ej: Volcanes, construir un modelo...")
            num_act_input_2 = st.number_input("Número de actividades en la secuencia", min_value=1, max_value=10, value=3, step=1, key="num_act_input_2")
            if st.button("Usar estas Ideas y Continuar", key="ideas_btn"):
                st.session_state.num_actividades = num_act_input_2
                handle_inspiration_submit(idea_box)
        with tab3:
            uploaded_file = st.file_uploader("Sube tu archivo .docx", type=['docx'])
            num_act_input_3 = st.number_input("Número de actividades en la secuencia", min_value=1, max_value=10, value=3, step=1, key="num_act_input_3")
            if st.button("Usar este Archivo y Continuar", key="file_btn"):
                if uploaded_file:
                    st.session_state.num_actividades = num_act_input_3
                    handle_inspiration_submit(leer_docx(uploaded_file))


    elif st.session_state.stage == "planning":
        st.header("ETAPA 2: Plan de Vuelo ✈️")
        st.markdown("Define los parámetros generales y deja que la IA diseñe una ruta de aprendizaje para la secuencia completa.")

        st.info("**Inspiración proporcionada:**")
        st.text_area("", value=st.session_state.inspiration_text, height=100, disabled=True)
        
        nivel_salida_final = st.selectbox("Máxima habilidad de Bloom a alcanzar AL FINAL de la secuencia", options=list(bloom_taxonomy_detallada.keys()), index=len(bloom_taxonomy_detallada) - 1)

        if st.button("🗺️ Generar Plan de Secuencia", type="primary"):
            with st.spinner("Creando el plan maestro..."):
                st.session_state.sequence_plan = planificar_secuencia(
                    st.session_state.inspiration_text, 
                    st.session_state.num_actividades,
                    nivel_salida_final,
                    st.session_state.gen_model_name
                )
        
        if st.session_state.sequence_plan:
            st.subheader("Plan de Secuencia Propuesto")
            st.markdown(st.session_state.sequence_plan)
            if st.button("✅ Me parece bien, ¡a generar las actividades!"):
                set_stage("generation")
                st.rerun()

        if st.button("Volver a Inspiración"):
            set_stage("inspiration")
            st.rerun()


    elif st.session_state.stage in ["generation", "display_sequence"]:
        st.header("ETAPA 3: Generación de la Secuencia 📚")
        
        with st.expander("Ver Plan de Secuencia Final", expanded=False):
            st.markdown(st.session_state.sequence_plan)
        
        col_p1, col_p2 = st.columns(2)
        with col_p1:
            st.subheader("Parámetros Pedagógicos")
            categoria_seleccionada = st.selectbox("Categoría", list(CATEGORIAS_ACTIVIDADES.keys()))
            if categoria_seleccionada == "Círculos de Matemática y Razonamiento":
                sub_options = CATEGORIAS_ACTIVIDADES[categoria_seleccionada]["Edades"]
            else:
                sub_options = CATEGORIAS_ACTIVIDADES[categoria_seleccionada]["Disciplinas"]
            subcategoria_seleccionada = st.selectbox("Grupo", sub_options)
        with col_p2:
            st.subheader("Nivel de Entrada")
            nivel_entrada_usuario = st.text_input("Nivel de entrada para la PRIMERA sesión", placeholder="Ej: Saben construir con bloques.")

        if st.button("🚀 Generar SECUENCIA COMPLETA con Auditoría", type="primary"):
            if not all([st.session_state.sequence_plan, nivel_entrada_usuario]):
                st.error("Por favor, asegúrate de tener un plan de secuencia y de definir el nivel de entrada.")
            else:
                with st.spinner(f"Generando y auditando las {st.session_state.num_actividades} actividades... Esto puede tardar."):
                    lista_actividades = []
                    
                    # Extraer los niveles de Bloom del plan
                    bloom_levels_per_session = re.findall(r"\*\*Objetivo Cognitivo \(Bloom\):\s*(\w+)", st.session_state.sequence_plan)

                    for i in range(1, st.session_state.num_actividades + 1):
                        params = {
                            "plan_secuencia": st.session_state.sequence_plan,
                            "session_num": i,
                            "grupo": subcategoria_seleccionada,
                            "nivel_entrada": nivel_entrada_usuario, # Se puede hacer más dinámico en el futuro
                            "nivel_salida": bloom_levels_per_session[i-1] if i-1 < len(bloom_levels_per_session) else "CREAR",
                            "gen_model": st.session_state.gen_model_name,
                            "audit_model": st.session_state.audit_model_name
                        }
                        actividad_generada = generar_actividad_con_auditoria(params)
                        lista_actividades.append(actividad_generada)
                
                st.session_state.processed_sequence = lista_actividades
                set_stage("display_sequence")
                st.rerun()

        if st.session_state.stage == "display_sequence" and st.session_state.processed_sequence:
            st.markdown("---")
            st.header("ETAPA 4: Secuencia de Aprendizaje Generada 🗺️")
            st.success("¡La secuencia completa ha sido generada!")

            for i, actividad_data in enumerate(st.session_state.processed_sequence):
                with st.expander(f"**Sesión {i+1}: {actividad_data.get('title', 'Sin Título')}** ({actividad_data.get('status', '❓')})"):
                    
                    texto_completo = actividad_data["activity_text"]
                    # Lógica para separar las guías
                    parts = re.split(r'---+\s*### GUÍA PARA EL DOCENTE', texto_completo, flags=re.IGNORECASE)
                    guia_rapida = parts[0].replace("### GUÍA RÁPIDA (PARA EL AULA / FICHA)", "").strip()
                    guia_docente = ""
                    if len(parts) > 1:
                        guia_docente = parts[1].strip()
                    else:
                        guia_rapida = texto_completo # Fallback si no encuentra el separador

                    tab1, tab2 = st.tabs(["Guía Rápida (Aula)", "Guía Docente (Acompañamiento)"])
                    with tab1:
                        st.markdown(guia_rapida)
                    with tab2:
                        st.markdown(guia_docente)

            st.markdown("---")
            st.subheader("✅ Exportar Secuencia Completa")
            word_buffer = exportar_secuencia_a_word(st.session_state.processed_sequence)
            st.download_button(
                label="Descargar Secuencia Completa en Word",
                data=word_buffer,
                file_name=f"secuencia_{subcategoria_seleccionada.replace(' ', '_')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        if st.session_state.stage in ["planning", "generation", "display_sequence"]:
            if st.button("Reiniciar y Empezar de Nuevo"):
                for key in list(st.session_state.keys()):
                    del st.session_state[key]
                st.rerun()

# --- Punto de Entrada del Script ---
if __name__ == "__main__":
    main()
